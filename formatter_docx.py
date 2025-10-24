"""Generate meeting minutes document using python-docx."""
from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document  # type: ignore
from docx.shared import Pt


TEMPLATE_PRESETS = {
    "通用": {
        "title_suffix": "会议纪要",
        "summary_heading": "议题要点",
        "summary_intro": None,
        "diff_heading": "录音校对差异",
        "action_heading": "行动项清单",
        "action_headers": ["责任人", "事项", "时间节点"],
        "policy_heading": "制度提示表",
        "policy_headers": ["制度标题", "条款", "来源", "摘要"],
    },
    "党委会": {
        "title_suffix": "党委会纪要",
        "summary_heading": "党委会议题概览",
        "summary_intro": "聚焦党委会重点议题，以下内容仅提示内部学习使用。",
        "diff_heading": "会议记录差异提醒",
        "action_heading": "整改/落实清单",
        "action_headers": ["责任部门", "整改事项", "时限要求"],
        "policy_heading": "党内制度提示",
        "policy_headers": ["制度名称", "条款号", "来源", "提示摘要"],
    },
    "项目会": {
        "title_suffix": "项目推进纪要",
        "summary_heading": "项目推进要点",
        "summary_intro": "以下为项目里程碑及关键讨论摘要。",
        "diff_heading": "会议记录核对",
        "action_heading": "项目行动项",
        "action_headers": ["责任人", "任务内容", "完成时间"],
        "policy_heading": "项目制度提示",
        "policy_headers": ["制度标题", "条款", "来源", "提示摘要"],
    },
    "招采会": {
        "title_suffix": "招采会议纪要",
        "summary_heading": "采购议题要点",
        "summary_intro": "重点关注合规流程与供应商议题，以下提示仅供参考。",
        "diff_heading": "录音核对差异",
        "action_heading": "采购执行清单",
        "action_headers": ["责任岗位", "执行事项", "完成时限"],
        "policy_heading": "招采制度提示",
        "policy_headers": ["制度名称", "条款", "来源", "提示摘要"],
    },
}


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    run.bold = True


def _add_kv_paragraph(document: Document, key: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run_key = paragraph.add_run(f"{key}：")
    run_key.bold = True
    paragraph.add_run(value)


def load_action_items(path: Path) -> List[Dict[str, str]]:
    if not path.exists():
        return []
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def create_minutes_document(
    output_path: Path,
    meeting_info: Dict[str, str],
    summary_title: str,
    summary_content: str,
    diff_content: Optional[str],
    action_items: List[Dict[str, str]],
    policy_suggestions: List[Dict[str, str]],
    template_name: str = "通用",
) -> Path:
    template = TEMPLATE_PRESETS.get(template_name, TEMPLATE_PRESETS["通用"])
    document = Document()
    document.styles["Normal"].font.size = Pt(11)

    base_title = meeting_info.get("title") or template["title_suffix"]
    if template_name != "通用" and template["title_suffix"] not in base_title:
        title_text = f"{base_title}（{template['title_suffix']}）"
    else:
        title_text = base_title
    document.add_heading(title_text, level=0)
) -> Path:
    document = Document()
    document.styles["Normal"].font.size = Pt(11)

    document.add_heading(meeting_info.get("title", "会议纪要"), level=0)

    base_fields = [
        ("会议主题", meeting_info.get("topic", "")),
        ("时间地点", meeting_info.get("time_place", "")),
        ("主持人", meeting_info.get("host", "")),
        ("记录人", meeting_info.get("recorder", "")),
        ("与会人员", meeting_info.get("participants", "")),
    ]
    for key, value in base_fields:
        _add_kv_paragraph(document, key, value)

    _add_heading(document, template["summary_heading"], level=1)
    if template.get("summary_intro"):
        document.add_paragraph(template["summary_intro"] or "")
    _add_heading(document, "议题要点", level=1)
    document.add_paragraph(summary_title, style=None)
    for line in summary_content.splitlines():
        if line.startswith("-"):
            document.add_paragraph(line[1:].strip(), style="List Bullet")
        elif line.strip():
            document.add_paragraph(line.strip())

    if diff_content:
        _add_heading(document, template["diff_heading"], level=1)
        _add_heading(document, "录音校对差异", level=1)
        for line in diff_content.splitlines():
            if line.startswith("##"):
                paragraph = document.add_paragraph()
                run = paragraph.add_run(line.replace("##", "").strip())
                run.bold = True
            elif line.startswith("-"):
                document.add_paragraph(line[1:].strip(), style="List Bullet")
            elif line.strip():
                document.add_paragraph(line.strip())

    _add_heading(document, template["action_heading"], level=1)
    if action_items:
        table = document.add_table(rows=1, cols=3)
        header_cells = table.rows[0].cells
        headers = template["action_headers"]
    _add_heading(document, "行动项清单", level=1)
    if action_items:
        table = document.add_table(rows=1, cols=3)
        header_cells = table.rows[0].cells
        headers = ["责任人", "事项", "时间节点"]
        for cell, text in zip(header_cells, headers):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.bold = True
        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("who", "")
            row_cells[1].text = item.get("what", "")
            row_cells[2].text = item.get("when", "")
    else:
        document.add_paragraph("（无）")

    _add_heading(document, template["policy_heading"], level=1)
    document.add_paragraph("以下提示仅供参考，不构成合规结论。")
    limited_policy = policy_suggestions[:20]
    if policy_suggestions and len(policy_suggestions) > len(limited_policy):
        document.add_paragraph("仅展示前 20 条制度匹配结果。")
    if limited_policy:
        table = document.add_table(rows=1, cols=4)
        header_cells = table.rows[0].cells
        headers = template["policy_headers"]
        document.add_paragraph("暂无行动项。")

    _add_heading(document, template["policy_heading"], level=1)
    _add_heading(document, "制度提示表", level=1)
    document.add_paragraph("以下提示仅供参考，不构成合规结论。")
    if policy_suggestions:
        table = document.add_table(rows=1, cols=4)
        header_cells = table.rows[0].cells
        headers = template["policy_headers"]
        headers = ["制度标题", "条款", "来源", "摘要"]
        for cell, text in zip(header_cells, headers):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.bold = True
        for suggestion in limited_policy:
        for suggestion in policy_suggestions:
            row = table.add_row().cells
            row[0].text = suggestion.get("title", "")
            row[1].text = suggestion.get("section", "")
            row[2].text = suggestion.get("source", "")
            row[3].text = suggestion.get("snippet", "")
    else:
        document.add_paragraph("未匹配到相关制度。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
